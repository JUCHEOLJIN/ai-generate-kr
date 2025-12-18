# korean-agent-project/app.py
import os
import json
from docx.enum.section import WD_SECTION
from dotenv import load_dotenv
from flask import Flask, make_response, request, jsonify, send_file
from google import genai
from google.genai import types
from flask_cors import CORS
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from io import BytesIO
from urllib.parse import quote  

# .env 파일에서 환경 변수를 로드
load_dotenv()

# --- 초기 설정 ---
app = Flask(__name__)

# CORS 설정
CORS(app, resources={r"/*": {"origins": ["http://localhost:5174", "https://ai-generate-kr-client.vercel.app"]}})

# LLM 클라이언트 초기화
try:
    # load_dotenv()로 로드된 API 키를 자동으로 사용
    client = genai.Client()
    print("✅ Gemini Client initialized successfully.")
except Exception as e:
    print(f"❌ Error initializing Gemini Client: {e}")
    client = None


# 필요할 때 주석 해제해서 사용.
print("--- 사용 가능한 모델 목록 ---")
try:
    # 모든 모델을 가져와서 이름만이라도 먼저 출력해 봅니다.
    models = client.models.list()
    for model in models:
        # 객체의 속성을 직접 확인하기 위해 __dict__ 또는 dir()을 참고할 수 있지만
        # 가장 기본이 되는 .name 속성은 반드시 존재합니다.
        print(f"모델 이름: {model.name}")
        
        # 상세 정보가 있다면 출력 (속성 존재 여부 체크)
        if hasattr(model, 'display_name'):
            print(f"표시 이름: {model.display_name}")
        
        # 지원 메서드는 아래와 같이 출력 시도
        if hasattr(model, 'supported_methods'):
            print(f"지원 메서드: {model.supported_methods}")
        
        print("-" * 40)
except Exception as e:
    print(f"오류 발생: {e}")


GEMINI_MODEL = 'models/gemini-flash-latest'


# --- 전략 선택 (Strategist) 정의 ---
# 장르별로 서로 다른 출제 가이드라인을 제공합니다.
STRATEGIES = {
    "문학_시": {
        "description": "운율, 비유, 상징, 화자의 태도 등 정서적/표현적 측면 분석",
        "points": "시어의 함축적 의미, 반어/역설 등 수사법, 화자의 정서 변화"
    },
    "문학_소설": {
        "description": "시점, 인물의 심리, 사건의 전개, 갈등 구조 분석",
        "points": "서술상의 특징(시점), 인물의 성격, 배경의 상징성"
    },
    "비문학": {
        "description": "정보의 사실적 이해, 논리적 추론, 핵심 개념 적용 분석",
        "points": "내용 일치, 전개 방식, <보기> 사례 적용, 어휘의 문맥적 의미"
    }
}


# --- 학년별 난이도 정의 ---
LEVEL_SPECIFIC_GUIDE = {
    "중1": "기초 문해력 단계. 지문의 사실 관계를 그대로 묻는 확인형 문제를 출제하세요.",
    "중2": "개념 이해 단계. 비유, 상징 등 기초적인 표현 기법의 특징을 묻는 문제를 포함하세요.",
    "중3": "중등 완성 단계. 문장 간의 관계를 파악하고 필자의 의도를 추론하는 문제를 출제하세요.",
    "고1": "공통 국어 단계. 내신 변별력을 위해 지문의 세부 사항을 꼼꼼하게 묻는 문제를 출제하세요.",
    "고2": "수능 준비 단계. 수능형 발문을 사용하고 개념을 사례에 적용하는 기초 추론을 넣으세요.",
    "고3": "수능 실전 단계. 복합적인 논리 구조를 파악하고 비판적으로 평가하는 고난도 문제를 생성하세요."
}


# --- 에이전트 코어 기능: 문제 생성 (현재는 단순 도구 역할) ---
def generate_problem_tool(input_text: str, num_problems: int, level="중1"):
    """에이전트 워크플로우: 지문 분석 후 맞춤형 문제 생성"""
    if not client:
        return {"error": "LLM Client not available."}

    try:
        # --- 1단계: Analyzer (지문 분석 및 장르 판별) ---
        analysis_prompt = f"""
        다음 국어 지문을 분석하여 [문학_시, 문학_소설, 비문학] 중 하나로 분류하고 이유를 설명하세요.
        반드시 다음 JSON 형식으로만 답하세요:
        {{"genre": "장르명", "reason": "분석 이유"}}

        [지문]:
        {input_text[:500]}...
        """
        
        analysis_response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=analysis_prompt,
            config=types.GenerateContentConfig(response_mime_type="application/json")
        )
        
        analysis_data = json.loads(analysis_response.text)
        genre = analysis_data.get("genre", "비문학")
        reason = analysis_data.get("reason", "")

        # --- 2단계: Strategist (전략 선택) ---
        # 판별된 장르가 STRATEGIES에 없으면 기본값으로 '비문학' 사용
        strategy = STRATEGIES.get(genre, STRATEGIES["비문학"])

        # --- 3단계: Leveler (학년별 난이도 선택) ---
        # 판별된 학년이 LEVEL_SPECIFIC_GUIDE에 없으면 기본값으로 '중1' 사용
        instruction = LEVEL_SPECIFIC_GUIDE.get(level, LEVEL_SPECIFIC_GUIDE["중1"])
        



        # --- 3단계: Generator (맞춤형 문제 생성) ---
        prompt = f"""
        당신은 대한민국 국어 교육 전문가입니다. 대상은 [{level}] 학생들입니다.
        분석 결과, 이 지문은 [{genre}] 타입입니다. ({reason})
        
        다음 지침에 따라 문제를 출제하십시오:
        1. 전략: {strategy['description']}
        2. 출제 포인트: {strategy['points']}
        3. 출제 대상 분석 지침: {instruction}
        4. 문제 개수: {num_problems}개
        5. 모든 문제는 5지선다형 객관식이어야 합니다.
        6. 2022 개정 교육과정의 성취기준을 바탕으로 하세요.

        출력은 반드시 제공된 JSON 스키마를 따르십시오.

        [지문]:
        {input_text}
        """

        # JSON 스키마 정의 (기존과 동일)
        json_schema = types.Schema(
            type=types.Type.OBJECT,
            properties={
                "problems": types.Schema(
                    type=types.Type.ARRAY,
                    items=types.Schema(
                        type=types.Type.OBJECT,
                        properties={
                            "type": types.Schema(type=types.Type.STRING),
                            "question": types.Schema(type=types.Type.STRING),
                            "options": types.Schema(type=types.Type.ARRAY, items=types.Schema(type=types.Type.STRING)),
                            "answer_index": types.Schema(type=types.Type.INTEGER),
                            "explanation": types.Schema(type=types.Type.STRING),
                        }
                    )
                )
            }
        )

        config = types.GenerateContentConfig(
            response_mime_type="application/json",
            response_schema=json_schema
        )

        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=prompt,
            config=config
        )

        # 결과 데이터 결합 (분석 결과도 함께 반환하여 프론트에서 활용 가능하게 함)
        result = json.loads(response.text)
        result["metadata"] = {
            "genre": genre,
            "analysis_reason": reason
        }

        return result

    except Exception as e:
        return {"error": f"Agent Workflow failed: {e}"}


# --- Flask API 엔드포인트 ---

@app.route('/generate', methods=['POST'])
def generate_endpoint():
    """사용자로부터 지문을 받아 문제를 생성하는 API"""
    data = request.json
    input_text = data.get('text', '').strip()
    num_problems = data.get('count', 1)
    level = data.get('level', '중1')

    if not input_text:
        return jsonify({"error": "Input text is required."}), 400

    # 현재는 단순 도구 호출
    result = generate_problem_tool(input_text, num_problems, level)

    # 에러가 있다면 500 응답, 아니면 200 응답
    if "error" in result:
        return jsonify(result), 500
    
    response_json = json.dumps(result, ensure_ascii=False, indent=4)
    response = app.response_class(
        response=response_json,
        status=200,
        mimetype='application/json'
    )
    return response


@app.route('/', methods=['GET'])
def health_check():
    """서버 상태 확인"""
    return "Korean Agent Project is Running!"



@app.route('/download-docx', methods=['POST'])
def download_docx():
    try:
        data = request.json
        problems = data.get('problems', [])
        level = data.get('level', '고등')
        
        doc = Document()

        # --- 1. 첫 번째 섹션 (제목 - 1단 기본값) ---
        section = doc.sections[0]
        # 여백 설정
        section.top_margin = Pt(40)
        section.bottom_margin = Pt(40)
        section.left_margin = Pt(40)
        section.right_margin = Pt(40)

        # 제목 추가 (현재 1단 상태)
        title = doc.add_heading(f'2025학년도 국어 시험지 ({level})', 0)
        title.alignment = 1 # 가운데 정렬
        doc.add_paragraph("") # 제목 아래 여백

        # --- 2. 두 번째 섹션 (문제 - 2단 설정) ---
        # 새로운 섹션을 추가합니다. WD_SECTION.CONTINUOUS는 페이지를 넘기지 않고 그 자리에서 섹션만 바꿉니다.
        new_section = doc.add_section(WD_SECTION.CONTINUOUS)
        
        # 새로운 섹션의 XML을 조작하여 2단으로 만듭니다.
        sectPr = new_section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), '2')  # 2단
        cols.set(qn('w:space'), '425')  # 단 사이 간격

        # 문제 루프 (이제부터 2단으로 들어갑니다)
        circles = ["①", "②", "③", "④", "⑤"]
        for i, p in enumerate(problems):
            para = doc.add_paragraph()
            run = para.add_run(f"{i+1}. {p['question']}")
            run.font.size = Pt(10.5)
            run.bold = True
            
            for idx, opt in enumerate(p['options']):
                opt_para = doc.add_paragraph()
                opt_para.paragraph_format.left_indent = Pt(10)
                opt_para.add_run(f"{circles[idx]} {opt}")
            
            doc.add_paragraph("")

        # --- 3. 세 번째 섹션 (정답지 - 다시 1단) ---
        doc.add_page_break() # 정답지는 다음 페이지로
        final_section = doc.add_section(WD_SECTION.NEW_PAGE)
        final_sectPr = final_section._sectPr
        final_cols = final_sectPr.xpath('./w:cols')[0]
        final_cols.set(qn('w:num'), '1') # 다시 1단

        doc.add_heading('정답 및 해설', level=1)
        for i, p in enumerate(problems):
            ans_para = doc.add_paragraph()
            ans_para.add_run(f"[{i+1}번 정답]: {p['answer_index'] + 1}번").bold = True
            ans_para.add_run(f"\n해설: {p['explanation']}")

        # 파일 저장 및 전송 로직 (기존과 동일)
        f = BytesIO()
        doc.save(f)
        binary_data = f.getvalue()
        f.close()
        
        response = make_response(binary_data)
        filename = f"ko_exam_{level}.docx"
        encoded_filename = quote(filename)

        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename={encoded_filename}'
        response.headers['Content-Length'] = len(binary_data)
        
        return response

    except Exception as e:
        print(f"다운로드 중 오류: {e}")
        return jsonify({"error": str(e)}), 500


# --- 서버 실행 ---
if __name__ == '__main__':
    # 개발 중에는 debug=True 설정
    app.run(debug=True, port=5000)